from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os

class WordProcessor:
    def __init__(self):
        pass
    
    def extract_content(self, file_path):
        """
        Extract content from a Word document (.docx or .doc)
        Returns a dictionary with text, images, and formatting information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            content = {
                'paragraphs': [],
                'tables': [],
                'images': [],
                'metadata': {
                    'title': '',
                    'author': '',
                    'subject': '',
                    'keywords': ''
                }
            }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_data = {
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment),
                        'runs': []
                    }
                    
                    # Extract run formatting
                    for run in paragraph.runs:
                        run_data = {
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_size': run.font.size.pt if run.font.size else None,
                            'font_name': run.font.name if run.font.name else None
                        }
                        para_data['runs'].append(run_data)
                    
                    content['paragraphs'].append(para_data)
            
            # Extract tables
            for table in doc.tables:
                table_data = {
                    'rows': []
                }
                
                for row in table.rows:
                    row_data = {
                        'cells': []
                    }
                    
                    for cell in row.cells:
                        cell_data = {
                            'text': cell.text,
                            'paragraphs': []
                        }
                        
                        for paragraph in cell.paragraphs:
                            cell_data['paragraphs'].append({
                                'text': paragraph.text,
                                'style': paragraph.style.name if paragraph.style else 'Normal'
                            })
                        
                        row_data['cells'].append(cell_data)
                    
                    table_data['rows'].append(row_data)
                
                content['tables'].append(table_data)
            
            # Extract metadata if available
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                content['metadata']['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                content['metadata']['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                content['metadata']['subject'] = doc.core_properties.subject
            if hasattr(doc.core_properties, 'keywords') and doc.core_properties.keywords:
                content['metadata']['keywords'] = doc.core_properties.keywords
            
            return content
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def get_document_info(self, file_path):
        """
        Get basic information about the Word document
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            info = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections),
                'file_size': os.path.getsize(file_path)
            }
            
            return info
            
        except Exception as e:
            raise Exception(f"Error getting document info: {str(e)}")
    
    def _remove_highlighting(self, run):
        """
        Remove highlighting from a run by both API and XML methods.
        This ensures highlighting is completely removed.
        """
        try:
            # Method 1: Use API
            if hasattr(run.font, 'highlight_color'):
                run.font.highlight_color = None
            # Method 2: Remove from XML directly - more thorough approach
            if hasattr(run, '_element'):
                rPr = run._element.get_or_add_rPr()
                if rPr is not None:
                    # Find all highlight elements (there might be multiple or nested)
                    namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                    # Method 1: Find by exact namespace
                    highlights = rPr.findall(f'{namespace}highlight')
                    for highlight in highlights:
                        rPr.remove(highlight)
                    # Method 2: Find by tag name (in case namespace is different)
                    for elem in list(rPr):
                        if elem.tag.endswith('}highlight') or 'highlight' in elem.tag.lower():
                            rPr.remove(elem)
                    # Method 3: Try to set the attribute directly if it exists
                    if hasattr(rPr, 'highlight'):
                        rPr.highlight = None
        except Exception:
            pass  # If highlighting can't be removed, continue
    
    def _remove_highlighting_from_all_runs(self, paragraph, start_idx, end_idx):
        """
        Remove highlighting from all runs between start_idx and end_idx (inclusive).
        This ensures we remove highlighting from all runs that were part of the placeholder.
        """
        for i in range(start_idx, min(end_idx + 1, len(paragraph.runs))):
            try:
                self._remove_highlighting(paragraph.runs[i])
            except Exception:
                pass
    
    def _remove_all_highlighting_from_paragraph(self, paragraph):
        """
        Remove highlighting from all runs in a paragraph.
        This is a comprehensive cleanup method to ensure no highlighting remains.
        """
        for run in paragraph.runs:
            try:
                self._remove_highlighting(run)
            except Exception:
                pass
    
    def _replace_placeholder_in_paragraph(self, paragraph, placeholder, value):
        """
        Replace a placeholder in a paragraph, handling cases where the placeholder
        may be split across multiple runs (e.g., due to formatting like bold).
        Preserves spacing and formatting around the placeholder.
        """
        placeholder_text = f'{{{placeholder}}}'
        
        # Check if placeholder exists in paragraph text
        if placeholder_text not in paragraph.text:
            return False
        
        # Convert value to string, handling None and empty values
        # Preserve the value as-is, including any whitespace it might have
        if value is None:
            replacement_value = ''
        elif isinstance(value, str) and value.lower() in ['nan', 'none', '']:
            replacement_value = ''
        else:
            replacement_value = str(value).strip()  # Strip only leading/trailing whitespace from value
        
        # Check if this is an email field - prevent line breaks in email addresses
        is_email_field = False
        if placeholder and 'email' in str(placeholder).lower():
            is_email_field = True
        elif replacement_value and '@' in replacement_value and '.' in replacement_value:
            # Also check if the value itself looks like an email
            is_email_field = True
        
        # For email addresses, prevent line breaks by using non-breaking characters
        if is_email_field and replacement_value:
            # Replace regular spaces with non-breaking spaces (Unicode \u00A0)
            replacement_value = replacement_value.replace(' ', '\u00A0')
            # Use zero-width non-breaking space (Unicode \u2060) after @ to prevent breaking there
            # This tells Word to treat the email as a single unbreakable unit at the @ symbol
            replacement_value = replacement_value.replace('@', '@\u2060')
        
        # First, try simple replacement if placeholder is in a single run
        for run in paragraph.runs:
            if placeholder_text in run.text:
                # Remove highlighting BEFORE replacement to ensure it's gone
                self._remove_highlighting(run)
                # Simple replacement - don't add extra spaces, just replace the placeholder
                # The template should already have correct spacing around placeholders
                run_text = run.text
                
                # Replace placeholder with value (no extra spacing)
                # If value is empty, just remove the placeholder
                if replacement_value == '':
                    replacement = ''
                else:
                    replacement = replacement_value
                
                # Replace the placeholder with the replacement value
                new_text = run_text.replace(placeholder_text, replacement)
                run.text = new_text
                
                # Remove highlighting again AFTER replacement to be thorough
                self._remove_highlighting(run)
                return True
        
        # Placeholder is split across multiple runs - need to handle this
        # Strategy: Replace text in the paragraph by working with runs
        full_text = paragraph.text
        placeholder_start = full_text.find(placeholder_text)
        
        if placeholder_start == -1:
            return False
        
        placeholder_end = placeholder_start + len(placeholder_text)
        
        # Find which runs contain parts of the placeholder
        current_pos = 0
        start_run_idx = None
        end_run_idx = None
        
        for i, run in enumerate(paragraph.runs):
            run_length = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_length
            
            if start_run_idx is None and run_start <= placeholder_start < run_end:
                start_run_idx = i
            if run_start < placeholder_end <= run_end:
                end_run_idx = i
                break
            
            current_pos = run_end
        
        if start_run_idx is None or end_run_idx is None:
            return False
        
        # Safety check: if start and end are the same, it should have been caught above
        # But handle it just in case
        if start_run_idx == end_run_idx:
            # Placeholder should be in a single run - try simple replacement
            if start_run_idx < len(paragraph.runs):
                run = paragraph.runs[start_run_idx]
                run_text = run.text
                placeholder_pos = run_text.find(placeholder_text)
                
                # Convert value to string, handling None and empty values
                if value is None:
                    replacement_value = ''
                elif isinstance(value, str) and value.lower() in ['nan', 'none', '']:
                    replacement_value = ''
                else:
                    replacement_value = str(value).strip()
                
                # Check if this is an email field - prevent line breaks in email addresses
                is_email_field = False
                if placeholder and 'email' in str(placeholder).lower():
                    is_email_field = True
                elif replacement_value and '@' in replacement_value and '.' in replacement_value:
                    is_email_field = True
                
                # For email addresses, prevent line breaks by using non-breaking characters
                if is_email_field and replacement_value:
                    replacement_value = replacement_value.replace(' ', '\u00A0')
                    replacement_value = replacement_value.replace('@', '@\u2060')
                
                # Simple replacement - don't add extra spaces
                # If value is empty, just remove the placeholder
                if replacement_value == '':
                    replacement = ''
                else:
                    replacement = replacement_value
                
                run.text = run_text.replace(placeholder_text, replacement)
                # Remove highlighting/background color
                self._remove_highlighting(run)
                return True
            return False
        
        # Calculate text before and after placeholder in the full text
        text_before = full_text[:placeholder_start]
        text_after = full_text[placeholder_end:]
        
        # Convert value to string, handling None and empty values (for multi-run case)
        if value is None:
            replacement_value = ''
        elif isinstance(value, str) and value.lower() in ['nan', 'none', '']:
            replacement_value = ''
        else:
            replacement_value = str(value).strip()
        
        # Check if this is an email field - prevent line breaks in email addresses
        is_email_field = False
        if placeholder and 'email' in str(placeholder).lower():
            is_email_field = True
        elif replacement_value and '@' in replacement_value and '.' in replacement_value:
            is_email_field = True
        
        # For email addresses, prevent line breaks by using non-breaking characters
        if is_email_field and replacement_value:
            replacement_value = replacement_value.replace(' ', '\u00A0')
            replacement_value = replacement_value.replace('@', '@\u2060')
        
        # Simple replacement for multi-run case - don't add extra spaces
        # If value is empty, just remove the placeholder
        if replacement_value == '':
            replacement = ''
        else:
            replacement = replacement_value
        
        # Now we need to reconstruct the paragraph
        # Save the original paragraph text with replacement
        new_paragraph_text = text_before + replacement + text_after
        
        # Calculate positions within individual runs
        pos = 0
        start_run_text_before = ''
        end_run_text_after = ''
        
        for i, run in enumerate(paragraph.runs):
            run_text = run.text
            run_len = len(run_text)
            
            if i == start_run_idx:
                # Calculate how much of this run is before the placeholder
                offset_in_run = placeholder_start - pos
                start_run_text_before = run_text[:offset_in_run]
            elif i == end_run_idx:
                # Calculate how much of this run is after the placeholder
                offset_in_run = placeholder_end - pos
                end_run_text_after = run_text[offset_in_run:]
                break
            
            pos += run_len
        
        # Remove highlighting from all runs that contained the placeholder BEFORE replacement
        # This ensures we catch highlighting that might be on any part of the placeholder
        self._remove_highlighting_from_all_runs(paragraph, start_run_idx, end_run_idx)
        
        # Replace the placeholder: update start run, remove middle runs, update/remove end run
        if start_run_idx < len(paragraph.runs):
            # Use the replacement value we calculated earlier (with spacing preserved)
            # Update start run with text before + replacement value
            paragraph.runs[start_run_idx].text = start_run_text_before + replacement
            # Remove highlighting/background color from the run (again, to be sure)
            self._remove_highlighting(paragraph.runs[start_run_idx])
        
        # Remove middle runs (between start and end, exclusive)
        runs_to_remove = list(range(start_run_idx + 1, end_run_idx))
        for i in reversed(runs_to_remove):
            if i < len(paragraph.runs):
                paragraph._element.remove(paragraph.runs[i]._element)
        
        # Handle end run - save formatting before removing runs
        orig_end_run_formatting = None
        orig_end_run_highlight = None
        if end_run_idx < len(paragraph.runs):
            orig_end_run = paragraph.runs[end_run_idx]
            orig_end_run_formatting = {
                'bold': orig_end_run.bold,
                'italic': orig_end_run.italic,
                'font_size': orig_end_run.font.size
            }
            orig_end_run_highlight = orig_end_run.font.highlight_color
        
        # Handle end run after removals
        if end_run_idx > start_run_idx:
            # After removing middle runs, the end_run_idx has shifted
            remaining_runs_count = len(paragraph.runs)
            expected_end_idx = end_run_idx - len(runs_to_remove)
            
            if expected_end_idx < remaining_runs_count and expected_end_idx > start_run_idx:
                # Update the end run with remaining text
                paragraph.runs[expected_end_idx].text = end_run_text_after
                # Remove highlighting if this run had part of the placeholder
                self._remove_highlighting(paragraph.runs[expected_end_idx])
            elif end_run_text_after:
                # Need to add a new run for the remaining text
                # Use formatting from the original end run if we saved it
                new_run = paragraph.add_run(end_run_text_after)
                if orig_end_run_formatting:
                    if orig_end_run_formatting['bold'] is not None:
                        new_run.bold = orig_end_run_formatting['bold']
                    if orig_end_run_formatting['italic'] is not None:
                        new_run.italic = orig_end_run_formatting['italic']
                    if orig_end_run_formatting['font_size'] is not None:
                        new_run.font.size = orig_end_run_formatting['font_size']
                # Ensure no highlighting on new run
                self._remove_highlighting(new_run)
        
        return True
    
    def _normalize_key(self, key):
        """
        Normalize a key for case-insensitive and whitespace-normalized matching.
        Converts to lowercase and normalizes whitespace (multiple spaces to single space, strip).
        """
        if not key:
            return ''
        return ' '.join(str(key).strip().lower().split())
    
    def _is_address_2_or_3_placeholder(self, placeholder):
        """
        Check if a placeholder is for Address 2 or Address 3.
        Handles various formats like 'Address 2', 'Address2', 'Address Line 2', etc.
        """
        if not placeholder:
            return False
        placeholder_lower = str(placeholder).lower().strip()
        # Check for address 2 or 3 in various formats
        address_patterns = [
            'address 2', 'address2', 'address line 2', 'addressline2',
            'address 3', 'address3', 'address line 3', 'addressline3',
            'addr 2', 'addr2', 'addr 3', 'addr3',
            'address2', 'address3'  # Without space
        ]
        return any(pattern in placeholder_lower for pattern in address_patterns) or \
               placeholder_lower in ['address 2', 'address2', 'address 3', 'address3', 
                                     'addr 2', 'addr2', 'addr 3', 'addr3',
                                     'address line 2', 'address line2', 'addressline2',
                                     'address line 3', 'address line3', 'addressline3']
    
    def _is_empty_value(self, value):
        """
        Check if a value is empty (None, empty string, 'nan', 'none', etc.)
        """
        if value is None:
            return True
        value_str = str(value).strip()
        return value_str == '' or value_str.lower() in ['nan', 'none', '']
    
    def _find_placeholder_matches(self, text, data):
        """
        Find all placeholders in text and return a mapping of placeholder -> data_key.
        Uses case-insensitive and whitespace-normalized matching.
        """
        import re
        placeholder_pattern = r'\{([^}]+)\}'
        matches = {}
        
        # Create normalized data mapping: normalized_key -> (original_key, value)
        normalized_data = {}
        for key, value in data.items():
            normalized_key = self._normalize_key(key)
            # If multiple keys normalize to the same value, prefer exact match, then first occurrence
            if normalized_key not in normalized_data:
                normalized_data[normalized_key] = (key, value)
            elif key == normalized_key:  # Prefer exact match
                normalized_data[normalized_key] = (key, value)
        
        # Find all placeholders in text
        found_placeholders = re.findall(placeholder_pattern, text)
        for placeholder in found_placeholders:
            normalized_placeholder = self._normalize_key(placeholder)
            if normalized_placeholder in normalized_data:
                original_key, value = normalized_data[normalized_placeholder]
                matches[placeholder] = (original_key, value)
        
        return matches
    
    def fill_placeholders(self, template_path, output_path, data):
        """
        Fill placeholders in the format {FieldName} in the Word template with values from data dict.
        Save the filled document to output_path.
        Handles placeholders that may be split across multiple runs due to formatting.
        Uses case-insensitive and whitespace-normalized matching for flexible column name matching.
        For trainee templates, removes empty paragraphs for Address 2 and Address 3 when they are empty.
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Check if this is a trainee template
        is_trainee_template = 'trainee' in os.path.basename(template_path).lower()
        
        try:
            doc = Document(template_path)
        except Exception as e:
            raise Exception(f"Error opening template: {str(e)}")
        
        try:
            # Track paragraphs that should be removed (for trainee template address 2/3)
            paragraphs_to_remove = []
            
            # Replace in paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # Store original text to check if paragraph only contains address placeholder
                original_text = paragraph.text
                
                # First, try exact matches (for backward compatibility and performance)
                for key, value in data.items():
                    try:
                        self._replace_placeholder_in_paragraph(paragraph, key, value)
                    except Exception as e:
                        # Log but continue processing
                        pass
                
                # Then, try case-insensitive and whitespace-normalized matches
                # This handles cases where Excel column names don't exactly match template placeholders
                # Check current paragraph text (after exact matches) for remaining placeholders
                try:
                    current_text = paragraph.text
                    matches = self._find_placeholder_matches(current_text, data)
                    for placeholder, (data_key, value) in matches.items():
                        # Only replace if placeholder still exists (exact match didn't work)
                        if f'{{{placeholder}}}' in current_text:
                            # Check if this is address 2 or address 3 for trainee template
                            is_address_2_or_3 = self._is_address_2_or_3_placeholder(placeholder)
                            is_empty = self._is_empty_value(value)
                            
                            # Replace the placeholder
                            self._replace_placeholder_in_paragraph(paragraph, placeholder, value)
                            
                            # For trainee template, if address 2 or 3 is empty, mark paragraph for removal
                            if is_trainee_template and is_address_2_or_3 and is_empty:
                                # Check if paragraph is now empty or only whitespace
                                updated_text = paragraph.text.strip()
                                if not updated_text or updated_text == '':
                                    paragraphs_to_remove.append(para_idx)
                            
                            # Update current_text for next iteration
                            current_text = paragraph.text
                except Exception as e:
                    # Log but continue processing
                    pass
                
                # Also check for address 2/3 in exact matches
                if is_trainee_template:
                    current_text_after = paragraph.text.strip()
                    # Check if this paragraph originally contained address 2 or 3 placeholder and is now empty
                    # Extract placeholders from original text using regex
                    import re
                    placeholder_pattern = r'\{([^}]+)\}'
                    original_placeholders = re.findall(placeholder_pattern, original_text)
                    has_address_2_3_placeholder = any(self._is_address_2_or_3_placeholder(p) for p in original_placeholders)
                    if has_address_2_3_placeholder and not current_text_after:
                        if para_idx not in paragraphs_to_remove:
                            paragraphs_to_remove.append(para_idx)
            
            # Remove empty paragraphs for address 2/3 in trainee template (in reverse order to maintain indices)
            if paragraphs_to_remove:
                # Get unique sorted indices in reverse order
                unique_indices = sorted(set(paragraphs_to_remove), reverse=True)
                for para_idx in unique_indices:
                    # Get paragraph element before removal (to avoid index issues)
                    if para_idx < len(doc.paragraphs):
                        para_element = doc.paragraphs[para_idx]._element
                        para_element.getparent().remove(para_element)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Track paragraphs to remove in this cell
                        cell_paragraphs_to_remove = []
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            original_text = paragraph.text
                            
                            # First, try exact matches
                            for key, value in data.items():
                                try:
                                    self._replace_placeholder_in_paragraph(paragraph, key, value)
                                except Exception as e:
                                    # Log but continue processing
                                    pass
                            
                            # Then, try case-insensitive and whitespace-normalized matches
                            # Check current paragraph text (after exact matches) for remaining placeholders
                            try:
                                current_text = paragraph.text
                                matches = self._find_placeholder_matches(current_text, data)
                                for placeholder, (data_key, value) in matches.items():
                                    # Only replace if placeholder still exists (exact match didn't work)
                                    if f'{{{placeholder}}}' in current_text:
                                        # Check if this is address 2 or address 3 for trainee template
                                        is_address_2_or_3 = self._is_address_2_or_3_placeholder(placeholder)
                                        is_empty = self._is_empty_value(value)
                                        
                                        # Replace the placeholder
                                        self._replace_placeholder_in_paragraph(paragraph, placeholder, value)
                                        
                                        # For trainee template, if address 2 or 3 is empty, mark paragraph for removal
                                        if is_trainee_template and is_address_2_or_3 and is_empty:
                                            # Check if paragraph is now empty or only whitespace
                                            updated_text = paragraph.text.strip()
                                            if not updated_text or updated_text == '':
                                                cell_paragraphs_to_remove.append(para_idx)
                                        
                                        # Update current_text for next iteration
                                        current_text = paragraph.text
                            except Exception as e:
                                # Log but continue processing
                                pass
                            
                            # Also check for address 2/3 in exact matches
                            if is_trainee_template:
                                current_text_after = paragraph.text.strip()
                                # Check if this paragraph originally contained address 2 or 3 placeholder and is now empty
                                # Extract placeholders from original text using regex
                                import re
                                placeholder_pattern = r'\{([^}]+)\}'
                                original_placeholders = re.findall(placeholder_pattern, original_text)
                                has_address_2_3_placeholder = any(self._is_address_2_or_3_placeholder(p) for p in original_placeholders)
                                if has_address_2_3_placeholder and not current_text_after:
                                    if para_idx not in cell_paragraphs_to_remove:
                                        cell_paragraphs_to_remove.append(para_idx)
                        
                        # Remove empty paragraphs in this cell (in reverse order)
                        if cell_paragraphs_to_remove:
                            # Get unique sorted indices in reverse order
                            unique_indices = sorted(set(cell_paragraphs_to_remove), reverse=True)
                            for para_idx in unique_indices:
                                # Get paragraph element before removal (to avoid index issues)
                                if para_idx < len(cell.paragraphs):
                                    para_element = cell.paragraphs[para_idx]._element
                                    para_element.getparent().remove(para_element)
            
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"Error filling placeholders: {str(e)}") 